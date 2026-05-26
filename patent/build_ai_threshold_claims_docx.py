from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "ai_threshold_switching_claims_draft.md"
SPEC_SOURCE = ROOT / "ai_threshold_switching_specification_draft.md"
OUT = ROOT / "ai_threshold_switching_claims_draft_v6.docx"


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def set_line_spacing(paragraph, line: float = 1.25) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(int(line * 240)))
    spacing.set(qn("w:lineRule"), "auto")


def parse_script_token(text: str, start: int) -> tuple[str, int]:
    if start >= len(text):
        return "", start
    if text[start] == "{":
        depth = 1
        idx = start + 1
        while idx < len(text) and depth > 0:
            if text[idx] == "{":
                depth += 1
            elif text[idx] == "}":
                depth -= 1
            idx += 1
        return text[start + 1 : idx - 1], idx
    if text[start] == "(":
        depth = 1
        idx = start + 1
        while idx < len(text) and depth > 0:
            if text[idx] == "(":
                depth += 1
            elif text[idx] == ")":
                depth -= 1
            idx += 1
        return text[start:idx], idx

    idx = start
    while idx < len(text):
        ch = text[idx]
        if ch.isspace() or ch in ",，。；;：:、=+-*/[]<>":
            break
        if ch in "_^":
            break
        idx += 1
    return text[start:idx], idx


def add_formula_like_runs(paragraph, text: str, bold: bool = False) -> None:
    idx = 0
    buffer = []

    def flush() -> None:
        nonlocal buffer
        if buffer:
            run = paragraph.add_run("".join(buffer))
            run.bold = bold
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(12)
            buffer = []

    while idx < len(text):
        ch = text[idx]
        if ch in "_^" and idx > 0:
            flush()
            token, next_idx = parse_script_token(text, idx + 1)
            if not token:
                buffer.append(ch)
                idx += 1
                continue
            run = paragraph.add_run(token)
            run.bold = bold
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10)
            if ch == "_":
                run.font.subscript = True
            else:
                run.font.superscript = True
            idx = next_idx
            continue
        buffer.append(ch)
        idx += 1
    flush()


def is_numbered_formula_line(text: str) -> bool:
    stripped = text.strip()
    if not re.search(r"（\d+）$", stripped):
        return False
    if "=" in stripped:
        return True
    if stripped.startswith("选择 "):
        return True
    return False


def normalize_formula_number_spacing(text: str) -> str:
    return re.sub(r"\s+（(\d+)）$", r"  （\1）", text)


def iter_blocks(lines: Iterable[str]) -> Iterable[list[str]]:
    block: list[str] = []
    for raw_line in lines:
        line = raw_line.rstrip("\n")
        if line.strip():
            block.append(line)
        else:
            if block:
                yield block
                block = []
    if block:
        yield block


def collect_markdown_blocks(text: str, stop_heading: str | None = None) -> list[list[str]]:
    blocks: list[list[str]] = []
    for block in iter_blocks(text.splitlines()):
        first = block[0].strip()
        if stop_heading is not None and first == stop_heading:
            break
        blocks.append(block)
    return blocks


def add_paragraph(doc: Document, text: str, *, bold: bool = False, center: bool = False, first_line: bool = True) -> None:
    text = normalize_formula_number_spacing(text)
    p = doc.add_paragraph()
    if is_numbered_formula_line(text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    if first_line and not center:
        if not is_numbered_formula_line(text):
            p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(3)
    add_formula_like_runs(p, text, bold=bold)
    set_line_spacing(p)


def add_claim(doc: Document, claim_no: str, claim_body_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{claim_no}. ")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    add_formula_like_runs(p, normalize_formula_number_spacing(claim_body_lines[0]), bold=False)
    set_line_spacing(p)

    for extra_line in claim_body_lines[1:]:
        sub = doc.add_paragraph()
        if is_numbered_formula_line(extra_line):
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            sub.paragraph_format.first_line_indent = Pt(24)
        sub.paragraph_format.space_after = Pt(3)
        add_formula_like_runs(sub, normalize_formula_number_spacing(extra_line), bold=False)
        set_line_spacing(sub)


def render_blocks(doc: Document, blocks: list[list[str]]) -> None:
    for block in blocks:
        first = block[0].strip()
        if first.startswith("# "):
            add_paragraph(doc, first[2:].strip(), bold=True, center=True, first_line=False)
            continue
        if first.startswith("## "):
            add_paragraph(doc, first[3:].strip(), bold=True, first_line=False)
            continue
        if first.startswith("### "):
            add_paragraph(doc, first[4:].strip(), bold=True, first_line=False)
            continue
        if first[:2].isdigit() and first[2:4] == ". ":
            claim_no = first.split(".", 1)[0]
            body_lines = [first.split(". ", 1)[1], *block[1:]]
            add_claim(doc, claim_no, body_lines)
            continue
        if len(first) >= 2 and first[0].isdigit() and first[1] == ".":
            claim_no = first.split(".", 1)[0]
            body_lines = [first.split(". ", 1)[1] if ". " in first else first.split(".", 1)[1].lstrip(), *block[1:]]
            add_claim(doc, claim_no, body_lines)
            continue
        if first.startswith("- "):
            for line in block:
                add_paragraph(doc, "• " + line[2:].strip(), first_line=False)
            continue
        for idx, line in enumerate(block):
            add_paragraph(doc, line, bold=False, center=False, first_line=(idx == 0))


def build_docx() -> Path:
    claims_text = SOURCE.read_text(encoding="utf-8")
    spec_text = SPEC_SOURCE.read_text(encoding="utf-8")
    doc = Document()
    set_document_style(doc)

    claims_blocks = collect_markdown_blocks(claims_text, stop_heading="## 说明")
    spec_blocks = collect_markdown_blocks(spec_text, stop_heading="## 写作建议")

    render_blocks(doc, claims_blocks)
    doc.add_page_break()
    render_blocks(doc, spec_blocks)

    doc.save(OUT)
    return OUT


def upgrade_equations_with_word_omml(doc_path: Path) -> None:
    import win32com.client

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    document = None
    try:
        document = word.Documents.Open(str(doc_path.resolve()))
        paragraph_count = document.Paragraphs.Count
        for index in range(1, paragraph_count + 1):
            paragraph = document.Paragraphs(index)
            text = paragraph.Range.Text.strip()
            if not text:
                continue
            if "=" not in text and not text.startswith("γ_work") and not text.startswith("s("):
                continue
            if "（" not in text and "。" not in text:
                continue
            if text.startswith("[") or text.startswith("其中") or text.startswith("构造") or text.startswith("当"):
                continue
            try:
                paragraph.Range.OMaths.Add(paragraph.Range)
                paragraph.Range.OMaths(1).BuildUp()
            except Exception:
                continue
        document.Save()
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


if __name__ == "__main__":
    out = build_docx()
    print(out)

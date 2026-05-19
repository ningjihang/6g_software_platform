from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "hybrid_mu_mimo_bicm_method_patent.docx"
FIG1 = ROOT / "fig1_hybrid_mu_mimo_bicm_flow.png"


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.bold = True


def set_line_spacing(paragraph, line: float = 1.25) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(int(line * 240)))
    spacing.set(qn("w:lineRule"), "auto")


def add_center(doc: Document, text: str, size: int = 16, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    set_line_spacing(p)


def add_para(doc: Document, text: str = "", first_line: bool = True) -> None:
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    set_line_spacing(p)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(16 if level == 1 else 13)
    set_line_spacing(p)


def add_claim(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"{number}. {text}")
    set_line_spacing(p)


def add_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        add_para(doc, step, first_line=False)


def add_image_center(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(caption)
    set_line_spacing(p)


def build() -> None:
    doc = Document()
    set_document_style(doc)

    add_para(doc, "申请人：复旦大学", first_line=False)
    add_para(doc, "发明人：________", first_line=False)
    add_para(doc, "联系人：________；手机：________；邮箱：________", first_line=False)
    add_para(doc, "申请文件要求提前公开：是", first_line=False)

    add_heading(doc, "说明书摘要")
    add_para(
        doc,
        "本发明属于无线通信技术领域，具体涉及一种面向多用户 MIMO-BICM 系统的混合预编码选择方法。"
        "该方法首先获取多用户下行信道矩阵及系统参数，并基于导频观测和由离线信道样本得到的信道均值及全协方差矩阵，"
        "对发射端信道状态信息进行全协方差 LMMSE 估计；随后根据估计信道构建满足恒模约束的共享模拟预编码矩阵，"
        "并在模拟预编码后的等效信道上为各用户构建块对角化零空间和降维等效信道；再生成包括 SVD、GMD、UCD 以及 GMD+THP 的结构化数字预编码候选；"
        "最后在相同符号样本和噪声样本下，对各混合预编码候选进行 BICM 广义互信息和误码率评价，并根据评价结果选择目标混合预编码矩阵。"
        "本发明能够在有限星座输入和非理想 CSI 条件下提高多用户 MIMO-BICM 系统的预编码选择可靠性，降低固定预编码结构在不同信噪比区间下的性能波动。",
    )

    add_heading(doc, "摘要附图")
    add_image_center(doc, FIG1, 14.5, "图1")

    doc.add_page_break()
    add_heading(doc, "权利要求书")

    add_claim(
        doc,
        1,
        "一种面向多用户 MIMO-BICM 系统的混合预编码选择方法，其特征在于，具体步骤如下：",
    )
    add_steps(
        doc,
        [
            "（1）获取多用户下行信道矩阵集合，并确定系统参数，所述系统参数包括发射天线数、接收天线数、用户数、射频链路数、每用户数据流数、调制阶数和信噪比；",
            "（2）基于所述下行信道矩阵集合获得发射端用于预编码设计的信道状态信息；",
            "（3）根据所述信道状态信息构建共享模拟预编码矩阵，所述共享模拟预编码矩阵满足恒模约束；",
            "（4）将各用户的下行信道矩阵与所述共享模拟预编码矩阵相乘，得到各用户的等效信道矩阵；",
            "（5）针对每个用户，根据其他用户的等效信道矩阵构建块对角化零空间，并将该用户的等效信道矩阵投影至所述块对角化零空间，得到该用户的降维等效信道矩阵；",
            "（6）基于所述降维等效信道矩阵生成多个结构化数字预编码候选，并将各用户的数字预编码候选组合成多用户数字预编码矩阵；",
            "（7）对由所述共享模拟预编码矩阵和所述多用户数字预编码矩阵形成的混合预编码候选进行接收端性能评价；",
            "（8）根据所述接收端性能评价结果，从所述混合预编码候选中选择目标混合预编码矩阵。",
        ],
    )

    add_claim(
        doc,
        2,
        "根据权利要求1所述的混合预编码选择方法，其特征在于，步骤（2）中，基于信道模型生成多个离线信道样本，"
        "对所述离线信道样本进行向量化处理，计算信道向量的样本均值和样本协方差矩阵，并对所述样本协方差矩阵施加对角加载，得到全协方差信道先验；"
        "基于重复单位矩阵导频获得导频观测，并利用所述全协方差信道先验对所述导频观测进行 LMMSE 估计，得到所述发射端用于预编码设计的信道状态信息。",
    )
    add_claim(
        doc,
        3,
        "根据权利要求2所述的混合预编码选择方法，其特征在于，所述重复单位矩阵导频的导频长度不小于发射天线数，且所述导频长度为发射天线数的整数倍。",
    )
    add_claim(
        doc,
        4,
        "根据权利要求1所述的混合预编码选择方法，其特征在于，步骤（3）中，分别对每个用户的信道状态信息进行奇异值分解，"
        "选取与每用户数据流数对应的右奇异向量，并提取所述右奇异向量的相位信息，由不同用户的相位信息构成满足恒模约束的共享模拟预编码矩阵。",
    )
    add_claim(
        doc,
        5,
        "根据权利要求1所述的混合预编码选择方法，其特征在于，步骤（5）中，对除目标用户之外的其他用户的等效信道矩阵进行堆叠，得到干扰信道矩阵；"
        "计算所述干扰信道矩阵的零空间基；将目标用户的等效信道矩阵与所述零空间基相乘，得到目标用户的降维等效信道矩阵。",
    )
    add_claim(
        doc,
        6,
        "根据权利要求1所述的混合预编码选择方法，其特征在于，步骤（6）中，所述结构化数字预编码候选至少包括："
        "基于所述降维等效信道矩阵奇异值分解得到的 SVD 数字预编码候选；基于几何均值分解目标对角增益得到的 GMD 数字预编码候选；"
        "基于增广奇异值谱和统一信道分解得到的 UCD 数字预编码候选；以及基于 GMD 数字预编码候选叠加 Tomlinson-Harashima 预编码的 THP 候选。",
    )
    add_claim(
        doc,
        7,
        "根据权利要求6所述的混合预编码选择方法，其特征在于，对于所述 THP 候选，在每个用户的上三角等效信道上，"
        "根据对角元素归一化后的上三角矩阵对发送符号进行逐层反馈抵消，并对反馈抵消后的符号执行模运算，得到 THP 发送符号。",
    )
    add_claim(
        doc,
        8,
        "根据权利要求7所述的混合预编码选择方法，其特征在于，步骤（7）中，对所述 THP 候选进行接收端性能评价包括："
        "对接收信号进行对角等化；对等化后的接收符号执行模折叠；枚举与所述模折叠对应的周期复制星座点；"
        "基于所述周期复制星座点计算各比特的对数似然比，并由所述对数似然比估计 BICM 广义互信息和误码率。",
    )
    add_claim(
        doc,
        9,
        "根据权利要求1所述的混合预编码选择方法，其特征在于，步骤（7）中，所述接收端性能评价在不同混合预编码候选之间使用相同的符号样本和噪声样本。",
    )
    add_claim(
        doc,
        10,
        "根据权利要求1所述的混合预编码选择方法，其特征在于，步骤（8）中，以多用户和速率最大为主评价指标，"
        "并结合用户间速率公平性、用户间泄漏功率与误码率中的至少一种指标，从所述混合预编码候选中选择目标混合预编码矩阵。",
    )

    doc.add_page_break()
    add_heading(doc, "说明书")
    add_center(doc, "一种面向多用户 MIMO-BICM 系统的混合预编码选择方法", size=15)

    add_heading(doc, "技术领域", level=2)
    add_para(doc, "[0001] 本发明属于无线通信技术领域，具体涉及一种面向多用户 MIMO-BICM 系统的混合预编码选择方法。")

    add_heading(doc, "背景技术", level=2)
    add_para(
        doc,
        "[0002] 多输入多输出系统能够利用多天线空间自由度提升无线通信系统的频谱效率。"
        "在毫米波通信、大规模天线阵列和多用户下行传输场景中，完全数字预编码需要为每根天线配置独立射频链路，硬件复杂度和功耗较高，"
        "因此由模拟预编码矩阵和数字预编码矩阵共同组成的混合预编码结构受到广泛关注。",
    )
    add_para(
        doc,
        "[0003] 在采用比特交织编码调制的 MIMO 系统中，实际传输通常使用有限阶 QAM 星座。"
        "与高斯输入容量优化不同，有限星座输入下的预编码设计不仅受信道奇异值影响，还受调制阶数、接收检测方式、误码传播和比特级互信息影响。",
    )
    add_para(
        doc,
        "[0004] 传统 SVD 预编码在低信噪比或强弱子信道差异明显时具有较好稳定性，GMD 或 UCD 类结构能够在特定高信噪比区间改善弱数据流性能，"
        "但在中等信噪比和非理想信道状态信息条件下可能出现性能波动。若仅根据发射端信道分解结果选择预编码矩阵，而不结合有限星座输入下的接收端 BICM 广义互信息评价，"
        "则可能导致所选预编码在实际接收链路中并非最优。",
    )

    add_heading(doc, "发明内容", level=2)
    add_para(
        doc,
        "[0005] 本发明的目的在于提供一种面向多用户 MIMO-BICM 系统的混合预编码选择方法，"
        "以解决有限星座输入和非理想 CSI 条件下固定预编码结构性能波动较大、预编码选择准则与接收端实际 BICM 性能不一致的问题。",
    )
    add_para(
        doc,
        "[0006] 为实现上述目的，本发明提供的混合预编码选择方法包括：获取多用户下行信道矩阵集合和系统参数；"
        "基于导频观测和全协方差信道先验获得发射端 CSI；根据所述 CSI 构建满足恒模约束的共享模拟预编码矩阵；"
        "对模拟预编码后的等效信道构建块对角化零空间和降维等效信道；生成结构化数字预编码候选；"
        "对混合预编码候选进行接收端 BICM 广义互信息和误码率评价；并根据评价结果选择目标混合预编码矩阵。",
    )
    add_para(
        doc,
        "[0007] 与现有技术相比，本发明至少具有如下有益效果：第一，将全协方差 LMMSE 信道估计引入混合预编码选择流程，"
        "能够利用信道相关性先验改善非理想 CSI 条件下的预编码可靠性；第二，在共享模拟预编码和多用户块对角化基础上同时构建 SVD、GMD、UCD 和 GMD+THP 等结构化候选，"
        "避免固定采用单一预编码结构；第三，对 THP 候选采用模折叠和周期复制星座的比特级评价方式，使预编码选择准则更贴近有限星座 MIMO-BICM 的实际接收性能。",
    )

    add_heading(doc, "附图说明", level=2)
    add_para(doc, "[0008] 图1为本发明面向多用户 MIMO-BICM 系统的混合预编码选择方法流程图。")

    add_heading(doc, "具体实施方式", level=2)
    add_para(doc, "[0009] 以下结合附图和实施例对本发明的技术方案进行详细说明。该实施例用于解释本发明，不构成对保护范围的限制。")
    add_para(
        doc,
        "[0010] 设系统包括 K 个用户，基站侧发射天线数为 N_t，射频链路数为 N_RF，每个用户接收天线数为 N_r，"
        "每个用户的数据流数为 d。第 k 个用户的信道矩阵为 H_k，发送端采用有限阶 QAM 星座并进行 BICM 传输。",
    )
    add_para(
        doc,
        "[0011] 在信道状态信息获取阶段，基于目标信道模型生成 N_cov 个信道样本，对每个样本 H^(n) 按列向量化为 h^(n)，"
        "计算样本均值和样本协方差矩阵，并对样本协方差矩阵施加对角加载。随后构建重复单位矩阵导频，基于导频观测和所述全协方差信道先验执行 LMMSE 信道估计，得到估计信道。",
    )
    add_para(
        doc,
        "[0012] 在共享模拟预编码构建阶段，对每个用户的估计信道执行奇异值分解，选取前 d 个右奇异向量并提取相位，"
        "将多个用户的相位向量组合为共享模拟预编码矩阵 F_RF，并使其满足 |[F_RF]_{i,j}|=1/sqrt(N_t) 的恒模约束。",
    )
    add_para(
        doc,
        "[0013] 在块对角化降维阶段，对目标用户 k，将其他用户的等效信道堆叠形成干扰信道矩阵，计算该干扰信道矩阵的零空间基 N_k，"
        "并得到目标用户降维等效信道 G_k N_k。该处理使每个用户的数字预编码块被限制在其他用户等效信道的零空间内，从而降低用户间干扰。",
    )
    add_para(
        doc,
        "[0014] 在结构化数字预编码候选生成阶段，对每个用户的降维等效信道执行奇异值分解。"
        "对于 SVD 候选，直接采用右奇异向量作为局部数字预编码基；对于 GMD 候选，根据奇异值计算几何均值目标对角增益，"
        "并通过右酉旋转将对角信道变换为具有目标对角增益的上三角等效信道；对于 UCD 候选，先构建增广奇异值谱，再根据统一信道分解目标生成局部数字预编码基。",
    )
    add_para(
        doc,
        "[0015] 对于 GMD+THP 候选，在每个用户的上三角等效信道上，先按其对角元素归一化得到反馈矩阵，"
        "再从后向前对发送符号进行反馈抵消和居中模运算，从而在发送端预先抵消用户内部上三角串扰。",
    )
    add_para(
        doc,
        "[0016] 在接收端评价阶段，对于线性候选，基于接收端投影后的信道进行并行检测或 SIC 检测，并计算 BICM 广义互信息和误码率。"
        "对于 THP 候选，接收端先进行对角等化和模折叠，再枚举周期复制星座点计算比特对数似然比，并由此估计 THP 感知 BICM 广义互信息和误码率。",
    )
    add_para(
        doc,
        "[0017] 为降低不同候选之间的随机比较误差，在一种实施方式中，不同混合预编码候选使用相同的符号样本和噪声样本进行 Monte Carlo 评价。"
        "最终选择使多用户和速率最大的候选作为目标混合预编码矩阵，也可以结合用户间公平性、用户间泄漏功率和误码率进行综合排序。",
    )
    add_para(
        doc,
        "[0018] 在一个仿真实施例中，用户数 K=2，发射天线数 N_t=16，每用户接收天线数 N_r=4，射频链路数 N_RF=8，"
        "每用户数据流数 d=4，调制方式采用 64QAM，信道模型采用 CDL-A，导频长度为 16，信道协方差样本数为 256，"
        "信噪比范围为 -10 dB 至 50 dB。可比较 SVD、GMD、GMD+THP、固定基线以及联合软优化后的混合预编码性能，并以 BICM 广义互信息和误码率作为评价指标。",
    )

    doc.add_page_break()
    add_heading(doc, "说明书附图")
    add_image_center(doc, FIG1, 15.5, "图1  多用户 MIMO-BICM 混合预编码选择方法流程图")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
